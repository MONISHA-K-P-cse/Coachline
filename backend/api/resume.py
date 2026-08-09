from fastapi import APIRouter, Depends, UploadFile, File, HTTPException, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from starlette.concurrency import run_in_threadpool
import pdfplumber
import docx
import logging
import io

from backend.core.database import get_db
from backend.core.auth import get_current_user
from backend.models.user import User
from backend.models.resume import Resume
from backend.schemas.resume import ResumeResponse

from ai.agents.resume_agent import ResumeAgent

router = APIRouter(prefix="/resume", tags=["Resume Pipeline"])
logger = logging.getLogger("resume")

resume_agent = ResumeAgent()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs_text = "\n".join([p.text for p in doc.paragraphs])
    
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_texts.append(" | ".join(row_text))
    
    if table_texts:
        paragraphs_text += "\n\n" + "\n".join(table_texts)
    return paragraphs_text


@router.post("/upload", response_model=ResumeResponse)
async def upload_resume(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    filename_lower = file.filename.lower()
    if not (filename_lower.endswith(".pdf") or filename_lower.endswith(".docx")):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only PDF and DOCX files are supported."
        )

    file_bytes = await file.read()
    extracted_text = ""

    try:
        if filename_lower.endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
        else:
            extracted_text = extract_text_from_docx(file_bytes)
    except Exception as exc:
        logger.error("Failed to parse resume: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Could not read this file - it may be corrupted or contain no extractable text.",
        )

    if not extracted_text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No text could be extracted from this PDF.",
        )

    # Real Resume Agent call (Granite LLM). ResumeAgent already has its own
    # internal fallback (fallback_used=True, score=50) if the model output
    # can't be parsed as structured JSON - this try/except only covers the
    # LLM call itself failing to run at all (e.g. Ollama not running).
    try:
        agent_score_data = await run_in_threadpool(resume_agent.analyze_resume, extracted_text)
    except Exception as exc:
        logger.warning("Resume agent call failed (%s); AI reviewer is unavailable.", exc)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="The AI resume reviewer is temporarily unavailable. Please try again shortly.",
        )

    resume_entry = Resume(
        user_id=current_user.id,
        filename=file.filename,
        parsed_text=extracted_text[:3000],
        score=agent_score_data["score"],
        ats_score=agent_score_data["ats_score"],
        keyword_count=agent_score_data["keyword_count"],
        resume_feedback=agent_score_data["resume_feedback"],
        score_details=agent_score_data,
    )

    db.add(resume_entry)
    db.commit()
    db.refresh(resume_entry)

    return ResumeResponse(
        id=resume_entry.id,
        filename=resume_entry.filename,
        score=resume_entry.score,
        ats_score=resume_entry.ats_score,
        keyword_count=resume_entry.keyword_count,
        resume_feedback=resume_entry.resume_feedback,
        parsed_text_preview=resume_entry.parsed_text[:200] if resume_entry.parsed_text else "",
        score_details=resume_entry.score_details,
        uploaded_at=resume_entry.uploaded_at
    )


@router.get("/", response_model=list[ResumeResponse])
def list_resumes(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    resumes = (
        db.query(Resume)
        .filter(Resume.user_id == current_user.id)
        .order_by(Resume.uploaded_at.desc())
        .all()
    )
    return [
        ResumeResponse(
            id=r.id,
            filename=r.filename,
            score=r.score,
            ats_score=r.ats_score,
            keyword_count=r.keyword_count,
            resume_feedback=r.resume_feedback,
            parsed_text_preview=r.parsed_text[:200] if r.parsed_text else "",
            score_details=r.score_details,
            uploaded_at=r.uploaded_at,
        )
        for r in resumes
    ]


@router.post("/{resume_id}/improve")
async def improve_resume_endpoint(
    resume_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    resume = db.query(Resume).filter(
        Resume.id == resume_id,
        Resume.user_id == current_user.id
    ).first()

    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    improvements = []
    if resume.score_details and "improvements" in resume.score_details:
        improvements = resume.score_details["improvements"]
    if not improvements:
        improvements = ["Add more quantifiable metrics", "Clarify cloud deployment tools"]

    try:
        opt_data = await run_in_threadpool(
            resume_agent.improve_resume,
            resume.parsed_text or "",
            improvements
        )
    except Exception as exc:
        logger.error("Failed to run resume optimizer: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="The resume optimizer is temporarily offline. Please try again later.",
        )

    return opt_data


class PDFGenerationRequest(BaseModel):
    text: str
    filename: str | None = "improved_resume.pdf"


@router.post("/generate-pdf")
async def generate_pdf_endpoint(req: PDFGenerationRequest):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF generation library is not available."
        )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    import re
    from reportlab.lib.enums import TA_CENTER

    def markdown_to_html(txt: str) -> str:
        escaped = txt.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        escaped = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', escaped)
        escaped = re.sub(r'\*(.*?)\*', r'<i>\1</i>', escaped)
        return escaped

    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        alignment=0, # Left-aligned
        textColor='#1C1917',
        spaceAfter=2
    )

    tagline_style = ParagraphStyle(
        name='TaglineStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11.5,
        leading=14.5,
        alignment=0,
        textColor='#D97706',
        spaceAfter=4
    )

    contact_style = ParagraphStyle(
        name='ContactStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11.5,
        alignment=0,
        textColor='#4B5563',
        spaceAfter=8
    )

    heading_style = ParagraphStyle(
        name='HeadingStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor='#1C1917',
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )

    date_style = ParagraphStyle(
        name='DateStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor='#1C1917'
    )

    role_style = ParagraphStyle(
        name='RoleStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=13.5,
        textColor='#1C1917'
    )

    dot_style = ParagraphStyle(
        name='DotStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=12,
        alignment=1, # Centered
        textColor='#D97706'
    )

    body_style = ParagraphStyle(
        name='BodyStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13.5,
        textColor='#374151',
        spaceAfter=2
    )

    skills_style = ParagraphStyle(
        name='SkillsStyle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=13.5,
        textColor='#4B5563',
        spaceBefore=4,
        spaceAfter=8
    )

    bullet_style = ParagraphStyle(
        name='BulletStyle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13.5,
        textColor='#374151',
        leftIndent=12,
        firstLineIndent=-12,
        spaceBefore=2,
        spaceAfter=2
    )

    story = []
    lines = req.text.split('\n')
    
    # 1. Parse Name, Tagline, and Contact Info
    name_str = ""
    tagline_str = ""
    contact_str = ""
    
    i = 0
    while i < len(lines):
        line_clean = lines[i].strip()
        if not line_clean:
            i += 1
            continue
        if line_clean.startswith('#'):
            name_str = line_clean.replace('#', '').strip()
            i += 1
            break
        elif len(line_clean) < 40 and not any(k in line_clean.lower() for k in ["@", "linkedin.com", "github.com", "phone", " | "]):
            name_str = line_clean
            i += 1
            break
        i += 1

    while i < len(lines):
        line_clean = lines[i].strip()
        if not line_clean:
            i += 1
            continue
        if not any(k in line_clean.lower() for k in ["@", "linkedin.com", "github.com", "phone", " | "]):
            tagline_str = line_clean
            i += 1
            break
        else:
            break

    while i < len(lines):
        line_clean = lines[i].strip()
        if not line_clean:
            i += 1
            continue
        if any(k in line_clean.lower() for k in ["@", "linkedin.com", "github.com", "phone", " | "]):
            contact_str = line_clean
            i += 1
            break
        else:
            break

    if name_str:
        story.append(Paragraph(name_str, title_style))
    if tagline_str:
        story.append(Paragraph(tagline_str, tagline_style))
    if contact_str:
        story.append(Paragraph(contact_str, contact_style))

    # Contact line divider bar
    story.append(Table([['']], colWidths=['100%'], rowHeights=[0.8], style=TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#D97706')),
        ('TOPPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
    ])))
    story.append(Spacer(1, 8))

    # 2. Parse sections
    sections = []
    current_sec = None

    while i < len(lines):
        line_clean = lines[i].rstrip()
        if not line_clean:
            i += 1
            continue
        content = line_clean.strip()

        is_header = False
        header_name = ""
        if content.startswith('##'):
            is_header = True
            header_name = content.replace('##', '').strip()
        elif content.isupper() and len(content) < 30 and not content.startswith(('-', '*', '•')):
            is_header = True
            header_name = content

        if is_header:
            current_sec = {
                "name": header_name,
                "type": "text",
                "items": []
            }
            name_lower = header_name.lower()
            if "skills" in name_lower:
                current_sec["type"] = "skills"
            elif any(k in name_lower for k in ["experience", "education", "volunteering"]):
                current_sec["type"] = "timeline"
            elif any(k in name_lower for k in ["strengths", "interests"]):
                current_sec["type"] = "grid"
            else:
                current_sec["type"] = "text"
            sections.append(current_sec)
            i += 1
            continue

        if not current_sec:
            i += 1
            continue

        if current_sec["type"] == "skills":
            skills = [s.strip() for s in content.split() if s.strip()]
            current_sec["items"].extend(skills)
            i += 1
        elif current_sec["type"] == "timeline":
            import re
            date_match = re.search(r'\b\d{2}/\d{2,4}\s*-\s*(\d{2}/\d{2,4}|Present|current)\b', content, re.I)
            if date_match:
                entry = {
                    "date": content,
                    "location": "",
                    "title": "",
                    "company": "",
                    "bullets": []
                }
                j = i + 1
                if j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(('-', '*', '•', '#')) and not re.search(r'\b\d{2}/\d{2,4}\b', lines[j]):
                    entry["location"] = lines[j].strip()
                    j += 1
                if j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(('-', '*', '•', '#')) and not re.search(r'\b\d{2}/\d{2,4}\b', lines[j]):
                    entry["title"] = lines[j].strip()
                    j += 1
                if j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(('-', '*', '•', '#')) and not re.search(r'\b\d{2}/\d{2,4}\b', lines[j]):
                    entry["company"] = lines[j].strip()
                    j += 1
                while j < len(lines):
                    next_line = lines[j].rstrip()
                    if not next_line:
                        j += 1
                        continue
                    next_content = next_line.strip()
                    if next_content.startswith(('-', '*', '•')):
                        entry["bullets"].append(next_content[1:].strip())
                        j += 1
                    elif next_content.startswith('#') or (next_content.isupper() and len(next_content) < 30) or re.search(r'\b\d{2}/\d{2,4}\s*-\s*(\d{2}/\d{2,4}|Present|current)\b', next_content, re.I):
                        break
                    else:
                        if entry["bullets"]:
                            entry["bullets"].append(next_content)
                        else:
                            entry["company"] = (entry["company"] + " " + next_content).strip()
                        j += 1
                current_sec["items"].append(entry)
                i = j
            else:
                current_sec["items"].append(content)
                i += 1
        elif current_sec["type"] == "grid":
            title = content
            desc = ""
            j = i + 1
            if j < len(lines) and lines[j].strip() and not lines[j].strip().startswith('#'):
                desc = lines[j].strip()
                i = j + 1
            else:
                i += 1
            current_sec["items"].append({"title": title, "desc": desc})
        else:
            current_sec["items"].append(content)
            i += 1

    # 3. Render structured sections
    for sec in sections:
        if not sec["items"]:
            continue
        story.append(Paragraph(sec["name"].upper(), heading_style))
        story.append(Spacer(1, 8))

        if sec["type"] == "skills":
            formatted_skills = " &nbsp;&nbsp;&bull;&nbsp;&nbsp; ".join([f"<font color='#D97706'>{s}</font>" for s in sec["items"]])
            story.append(Paragraph(formatted_skills, skills_style))
            story.append(Spacer(1, 10))
        elif sec["type"] == "timeline":
            for entry in sec["items"]:
                if isinstance(entry, dict):
                    col1_text = f"<b>{entry['date']}</b>"
                    if entry['location']:
                        col1_text += f"<br/><font color='#6B7280'>{entry['location']}</font>"
                    col2_text = "<font color='#D97706'>&#9679;</font>"
                    
                    # Col 3 flowables list
                    col3_flowables = []
                    title_company_html = ""
                    if entry['title']:
                        title_company_html += f"<b>{entry['title']}</b>"
                    if entry['company']:
                        if title_company_html:
                            title_company_html += "<br/>"
                        title_company_html += f"<font color='#D97706'><b>{entry['company']}</b></font>"
                    
                    if title_company_html:
                        col3_flowables.append(Paragraph(title_company_html, role_style))
                        col3_flowables.append(Spacer(1, 2))
                    
                    for b in entry['bullets']:
                        bullet_html = f"<font color='#D97706'>&#9679;</font>&nbsp;&nbsp;{markdown_to_html(b)}"
                        col3_flowables.append(Paragraph(bullet_html, bullet_style))
                    
                    p1 = Paragraph(col1_text, date_style)
                    p2 = Paragraph(col2_text, dot_style)
                    
                    t = Table([[p1, p2, col3_flowables]], colWidths=[110, 20, 370])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('TOPPADDING', (0,0), (-1,-1), 6),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                    ]))
                    story.append(t)
                else:
                    story.append(Paragraph(markdown_to_html(entry), body_style))
            story.append(Spacer(1, 10))
        elif sec["type"] == "grid":
            grid_data = []
            row = []
            for item in sec["items"]:
                cell_text = f"<b>{item['title']}</b>"
                if item['desc']:
                    cell_text += f"<br/><font color='#4B5563'>{markdown_to_html(item['desc'])}</font>"
                p = Paragraph(cell_text, body_style)
                row.append(p)
                if len(row) == 2:
                    grid_data.append(row)
                    row = []
            if row:
                row.append('')
                grid_data.append(row)
            
            t = Table(grid_data, colWidths=[250, 250])
            t.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 10),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
        else:
            for txt in sec["items"]:
                if txt.startswith(('-', '*', '•')):
                    bullet_text = txt[1:].strip()
                    formatted_text = f'&nbsp;&nbsp;<font color="#D97706">&#9679;</font>&nbsp;&nbsp;{markdown_to_html(bullet_text)}'
                    story.append(Paragraph(formatted_text, body_style))
                else:
                    story.append(Paragraph(markdown_to_html(txt), body_style))
            story.append(Spacer(1, 10))
            
        story.append(Spacer(1, 14))

    doc.build(story)
    buffer.seek(0)

    filename = req.filename or "improved_resume.pdf"
    if not filename.endswith(".pdf"):
        filename += ".pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
